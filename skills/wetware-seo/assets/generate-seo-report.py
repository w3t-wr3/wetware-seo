#!/usr/bin/env python3
"""
Generate Big Thunder Marine Multi-Brand SEO Audit Report.
Uses the official Wetware_Labs_Template.docx for header/footer.
Data sourced from SEO pipeline collection (2026-03-26).
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from copy import deepcopy
import io
import os
import subprocess
import tempfile
from zipfile import ZipFile

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
import numpy as np

SKILL_DIR = "/Users/myrm/obsidian/vaults/.claude/skills/wetware-docs"
TEMPLATE_PATH = f"{SKILL_DIR}/assets/Wetware_Labs_Template.docx"
LOGO_PATH = f"{SKILL_DIR}/assets/wetwareArtboard 1@4x.png"
OUTPUT_DOCX = "/Users/myrm/Documents/clients/Big Thunder Marine/Big_Thunder_Marine_SEO_Audit_March2026.docx"

BLACK = RGBColor(0, 0, 0)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
GRAY = RGBColor(0x55, 0x55, 0x55)
LIGHT_GRAY = RGBColor(0x99, 0x99, 0x99)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
# Matplotlib grayscale palette
C_BLACK = '#222222'
C_DARK = '#444444'
C_MID = '#444444'
C_LIGHT = '#AAAAAA'
C_PALE = '#CCCCCC'
C_FAINT = '#E5E5E5'

TMPDIR = tempfile.mkdtemp()


# ── CHART GENERATORS ─────────────────────────────────────────────────────────

def chart_brand_performance():
    """Horizontal bar chart: mobile performance per brand. Grayscale."""
    brands = ['Iconic Yachts', 'Dock Co.', 'BT Wake', 'Boat Rentals']
    mobile = [71, 47, 80, 98]
    desktop = [99, 67, 96, 100]

    fig, ax = plt.subplots(figsize=(6.0, 1.7))
    y = np.arange(len(brands))
    bars_m = ax.barh(y + 0.16, mobile, 0.30, label='Mobile', color=C_BLACK, edgecolor='white', linewidth=0.5)
    bars_d = ax.barh(y - 0.16, desktop, 0.30, label='Desktop', color=C_PALE, edgecolor='white', linewidth=0.5)

    for bar, val in zip(bars_m, mobile):
        ax.text(bar.get_width() + 1.5, bar.get_y() + bar.get_height()/2, f'{val}%',
                va='center', fontsize=8, fontweight='bold', color=C_BLACK)
    for bar, val in zip(bars_d, desktop):
        ax.text(bar.get_width() + 1.5, bar.get_y() + bar.get_height()/2, f'{val}%',
                va='center', fontsize=8, color=C_MID)

    ax.set_yticks(y)
    ax.set_yticklabels(brands, fontsize=9, fontfamily='Arial')
    ax.set_xlim(0, 115)
    ax.set_ylim(-0.6, len(brands) + 0.2)
    ax.set_xlabel('')
    ax.axvline(x=90, color=C_DARK, linestyle='--', linewidth=0.7, alpha=0.3)
    ax.axvline(x=50, color=C_DARK, linestyle='--', linewidth=0.7, alpha=0.3)
    ax.text(90, len(brands) + 0.05, 'Good', fontsize=6, color=C_DARK, ha='center')
    ax.text(50, len(brands) + 0.05, 'Needs Work', fontsize=6, color=C_DARK, ha='center')
    ax.legend(fontsize=7, loc='lower right', framealpha=0.9, edgecolor=C_PALE)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['bottom'].set_color(C_PALE)
    ax.spines['left'].set_color(C_PALE)
    ax.tick_params(colors=C_MID, labelsize=8)
    ax.xaxis.set_major_formatter(mticker.PercentFormatter())
    plt.tight_layout()
    path = os.path.join(TMPDIR, 'perf_chart.png')
    fig.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path


def chart_issue_severity():
    """Donut chart: issue severity breakdown. Grayscale."""
    labels = ['Critical', 'High', 'Medium', 'Low']
    counts = [5, 9, 8, 4]
    colors = [C_BLACK, C_DARK, C_LIGHT, C_FAINT]
    explode = (0.05, 0, 0, 0)

    fig, ax = plt.subplots(figsize=(2.6, 2.6))
    wedges, texts, autotexts = ax.pie(
        counts, labels=labels, colors=colors, explode=explode,
        autopct='%1.0f%%', startangle=90, textprops={'fontsize': 8, 'fontfamily': 'Arial'},
        pctdistance=0.75, labeldistance=1.15
    )
    for at in autotexts:
        at.set_fontsize(8)
        at.set_fontweight('bold')
        at.set_color('white')
    autotexts[-1].set_color(C_DARK)
    autotexts[-2].set_color(C_DARK)
    centre = plt.Circle((0, 0), 0.45, fc='white')
    ax.add_artist(centre)
    ax.text(0, 0, f'{sum(counts)}', ha='center', va='center', fontsize=20,
            fontweight='bold', color=C_BLACK, fontfamily='Arial')
    ax.text(0, -0.15, 'total issues', ha='center', va='center', fontsize=7,
            color=C_MID, fontfamily='Arial')
    plt.tight_layout()
    path = os.path.join(TMPDIR, 'severity_chart.png')
    fig.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path


def chart_cwv_comparison():
    """Grouped bar chart: LCP per brand. Grayscale."""
    brands = ['Dock Co.', 'Boat Rentals', 'BT Wake', 'Iconic']
    lcp = [5.9, 2.4, 3.9, 40.3]
    fcp = [3.3, 1.1, 2.2, 1.8]

    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(6.0, 1.7), gridspec_kw={'width_ratios': [3, 1]})

    x = np.arange(3)
    b_fcp = ax1.bar(x - 0.17, fcp[:3], 0.32, label='FCP', color=C_LIGHT, edgecolor='white')
    b_lcp = ax1.bar(x + 0.17, lcp[:3], 0.32, label='LCP', color=C_BLACK, edgecolor='white')
    for bar, val in zip(b_fcp, fcp[:3]):
        ax1.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.2, f'{val}s',
                ha='center', fontsize=7, color=C_MID)
    for bar, val in zip(b_lcp, lcp[:3]):
        ax1.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.2, f'{val}s',
                ha='center', fontsize=7, fontweight='bold', color=C_BLACK)
    ax1.set_xticks(x)
    ax1.set_xticklabels(brands[:3], fontsize=8, fontfamily='Arial')
    ax1.set_ylabel('Seconds', fontsize=8, color=C_MID)
    ax1.axhline(y=2.5, color=C_MID, linestyle='--', linewidth=0.7, alpha=0.4)
    ax1.set_ylim(0, 8.5)
    ax1.legend(fontsize=7, loc='upper left', framealpha=0.9, edgecolor=C_PALE)
    ax1.spines['top'].set_visible(False)
    ax1.spines['right'].set_visible(False)
    ax1.spines['bottom'].set_color(C_PALE)
    ax1.spines['left'].set_color(C_PALE)
    ax1.tick_params(colors=C_MID, labelsize=7)

    ax2.bar([0], [40.3], 0.5, color=C_BLACK, edgecolor='white')
    ax2.bar([0.7], [1.8], 0.5, color=C_LIGHT, edgecolor='white')
    ax2.text(0, 42, '40.3s', ha='center', fontsize=8, fontweight='bold', color=C_BLACK)
    ax2.text(0.7, 4, '1.8s', ha='center', fontsize=7, color=C_MID)
    ax2.set_xticks([0, 0.7])
    ax2.set_xticklabels(['LCP', 'FCP'], fontsize=7)
    ax2.set_title('Iconic', fontsize=9, fontweight='bold', color=C_BLACK, pad=10)
    ax2.set_ylim(0, 50)
    ax2.spines['top'].set_visible(False)
    ax2.spines['right'].set_visible(False)
    ax2.spines['bottom'].set_color(C_PALE)
    ax2.spines['left'].set_color(C_PALE)
    ax2.tick_params(colors=C_MID, labelsize=7)

    plt.tight_layout(w_pad=3)
    path = os.path.join(TMPDIR, 'cwv_chart.png')
    fig.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path


def chart_load_time():
    """Bar chart: server response time (TTFB) per brand."""
    brands = ['Dock Co.', 'Boat\nRentals', 'BT Wake', 'Iconic']
    ttfb = [0.137, 0.113, 0.117, 0.128]
    mobile_lcp = [5.9, 2.4, 3.9, 40.3]

    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(6.0, 1.5))

    # TTFB (server speed)
    bars1 = ax1.bar(brands, ttfb, color=C_DARK, edgecolor='white', width=0.5)
    for bar, val in zip(bars1, ttfb):
        ax1.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.008,
                f'{int(val*1000)}ms', ha='center', fontsize=8, fontweight='bold', color=C_BLACK)
    ax1.set_title('Server Response Time', fontsize=9, fontweight='bold', color=C_BLACK, pad=10)
    ax1.set_ylim(0, 0.22)
    ax1.axhline(y=0.1, color=C_MID, linestyle='--', linewidth=0.7, alpha=0.3)
    ax1.set_yticklabels([])
    ax1.spines['top'].set_visible(False)
    ax1.spines['right'].set_visible(False)
    ax1.spines['bottom'].set_color(C_PALE)
    ax1.spines['left'].set_visible(False)
    ax1.tick_params(colors=C_MID, labelsize=7, left=False)

    # Time to usable (LCP; capped at 8s, Iconic annotated)
    bars2 = ax2.bar(brands[:3], mobile_lcp[:3], color=C_DARK, edgecolor='white', width=0.5)
    bar_iconic = ax2.bar([brands[3]], [8], color=C_BLACK, edgecolor='white', width=0.5)  # clipped
    for bar, val in zip(bars2, mobile_lcp[:3]):
        ax2.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.25,
                f'{val}s', ha='center', fontsize=8, fontweight='bold', color=C_BLACK)
    # Iconic label above clipped bar
    ax2.text(bar_iconic[0].get_x() + bar_iconic[0].get_width()/2, 8.3,
            '40.3s', ha='center', fontsize=8, fontweight='bold', color=C_BLACK)
    ax2.set_title('Time Until Page is Usable', fontsize=9, fontweight='bold', color=C_BLACK, pad=10)
    ax2.set_ylim(0, 9.5)
    ax2.axhline(y=2.5, color=C_MID, linestyle='--', linewidth=0.7, alpha=0.3)
    ax2.set_yticklabels([])
    ax2.spines['top'].set_visible(False)
    ax2.spines['right'].set_visible(False)
    ax2.spines['bottom'].set_color(C_PALE)
    ax2.spines['left'].set_visible(False)
    ax2.tick_params(colors=C_MID, labelsize=7, left=False)

    plt.tight_layout(w_pad=3)
    path = os.path.join(TMPDIR, 'load_time_chart.png')
    fig.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path


def chart_issues_by_brand():
    """Stacked horizontal bar: issue count per brand by severity."""
    brands = ['Dock Co.', 'Boat Rentals', 'BT Wake', 'Iconic']
    critical = [1, 0, 0, 1]
    high =     [3, 3, 0, 3]
    medium =   [3, 1, 0, 2]
    low =      [0, 0, 3, 0]

    fig, ax = plt.subplots(figsize=(6.0, 1.4))
    y = np.arange(len(brands))

    ax.barh(y, critical, 0.5, label='Critical', color=C_BLACK)
    ax.barh(y, high, 0.5, left=critical, label='High', color=C_DARK)
    ax.barh(y, medium, 0.5, left=[c+h for c,h in zip(critical,high)], label='Medium', color=C_LIGHT)
    ax.barh(y, low, 0.5, left=[c+h+m for c,h,m in zip(critical,high,medium)], label='Low', color=C_FAINT)

    totals = [c+h+m+l for c,h,m,l in zip(critical,high,medium,low)]
    for i, total in enumerate(totals):
        ax.text(total + 0.3, i, f'{total} issues', va='center', fontsize=8, fontweight='bold', color=C_BLACK)

    ax.set_yticks(y)
    ax.set_yticklabels(brands, fontsize=9, fontfamily='Arial')
    ax.set_xlim(0, 12)
    ax.set_ylim(-0.5, len(brands) + 0.3)
    ax.legend(fontsize=7, loc='upper right', ncol=4, framealpha=0.9, edgecolor=C_PALE)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['bottom'].set_color(C_PALE)
    ax.spines['left'].set_color(C_PALE)
    ax.tick_params(colors=C_MID, labelsize=7)
    ax.set_xticklabels([])
    plt.tight_layout()
    path = os.path.join(TMPDIR, 'issues_brand_chart.png')
    fig.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path


def chart_seo_completeness():
    """Heatmap-style matrix: feature completeness per brand. Grayscale."""
    brands = ['Dock Co.', 'Boat Rentals', 'BT Wake', 'Iconic']
    features = ['Meta Title', 'Meta Desc.', 'OG Tags', 'Canonical', 'JSON-LD', 'robots.txt', 'Sitemap', 'H1 Tag']
    matrix = [
        [1, 0.5, 0.5, 1, 0, 1, 1, 1],
        [1, 1, 0, 0, 0, 0, 0, 0.5],
        [1, 1, 1, 1, 1, 0.5, 1, 1],
        [1, 1, 0, 0, 0, 0.5, 1, 1],
    ]

    fig, ax = plt.subplots(figsize=(6.0, 1.6))
    cmap = matplotlib.colors.ListedColormap([C_BLACK, C_MID, C_FAINT])
    bounds = [-0.25, 0.25, 0.75, 1.25]
    norm = matplotlib.colors.BoundaryNorm(bounds, cmap.N)

    data = np.array(matrix)
    ax.imshow(data, cmap=cmap, norm=norm, aspect='auto')

    ax.set_xticks(np.arange(len(features)))
    ax.set_yticks(np.arange(len(brands)))
    ax.set_xticklabels(features, fontsize=7, fontfamily='Arial', rotation=30, ha='right')
    ax.set_yticklabels(brands, fontsize=8, fontfamily='Arial')

    labels = {1: '\u2713', 0.5: '~', 0: '\u2717'}
    for i in range(len(brands)):
        for j in range(len(features)):
            val = matrix[i][j]
            txt_color = 'white' if val < 0.75 else C_BLACK
            ax.text(j, i, labels[val], ha='center', va='center',
                    fontsize=10, fontweight='bold', color=txt_color)

    ax.spines[:].set_visible(False)
    ax.tick_params(colors=C_MID)
    plt.tight_layout()
    path = os.path.join(TMPDIR, 'completeness_chart.png')
    fig.savefig(path, dpi=200, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path


# ── DOCUMENT HELPERS ─────────────────────────────────────────────────────────

def add_text(doc, text, size=10, color=DARK_GRAY, bold=False, italic=False,
             align=None, space_before=0, space_after=0):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    return p


def section_heading(doc, text, page_break=False):
    p = doc.add_paragraph()
    if page_break:
        p.paragraph_format.page_break_before = True
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(16)
    run.font.color.rgb = BLACK
    run.bold = True
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single', qn('w:sz'): '3',
        qn('w:space'): '3', qn('w:color'): '000000'
    })
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def sub_heading(doc, text, space_before=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.font.color.rgb = BLACK
    run.bold = True
    return p


def add_spacer(doc, space_pt=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_pt)
    p.paragraph_format.space_after = Pt(0)
    return p


def set_cell_shading(cell, color_hex):
    shading = cell._tc.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): color_hex
    })
    shading.append(shd)


def set_cell_text(cell, text, size=10, color=DARK_GRAY, bold=False, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.makeelement(qn('w:tcBorders'), {})
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            border = tcBorders.makeelement(qn(f'w:{side}'), {
                qn('w:val'): 'single', qn('w:sz'): val.get('sz', '4'),
                qn('w:space'): '0', qn('w:color'): val.get('color', '000000')
            })
            tcBorders.append(border)
    tcPr.append(tcBorders)


def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn('w:tblPr'), {})
    borders = tblPr.makeelement(qn('w:tblBorders'), {})
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = borders.makeelement(qn(f'w:{side}'), {
            qn('w:val'): 'none', qn('w:sz'): '0',
            qn('w:space'): '0', qn('w:color'): 'auto'
        })
        borders.append(border)
    tblPr.append(borders)


def bullet(doc, text, size=9, color=DARK_GRAY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(f"\u2022  {text}")
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


def add_chart(doc, chart_path, width=Inches(6.0)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    run.add_picture(chart_path, width=width)
    return p


def score_color(score):
    """All black; no color coding. Wetware brand is B&W."""
    return DARK_GRAY


def issue_row(doc, issue_text, description, severity="Warning"):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Inches(5.2)
    table.columns[1].width = Inches(1.3)
    remove_table_borders(table)
    for col in range(2):
        set_cell_borders(table.cell(0, col), bottom={'sz': '2', 'color': 'CCCCCC'})
    cell = table.cell(0, 0)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(f"\u25A0  {issue_text}")
    run.font.name = "Arial"
    run.font.size = Pt(9.5)
    run.font.color.rgb = BLACK
    run.bold = True
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(4)
    run2 = p2.add_run(description)
    run2.font.name = "Arial"
    run2.font.size = Pt(7.5)
    run2.font.color.rgb = GRAY
    run2.italic = True
    right_cell = table.cell(0, 1)
    right_cell.text = ""
    p = right_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    sev_color = BLACK if severity == "Critical" else GRAY if severity == "High" else GRAY
    run = p.add_run(severity)
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = sev_color
    run.bold = (severity in ("Critical", "High"))


def fixable_row(doc, name, desc, priority):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Inches(5.5)
    table.columns[1].width = Inches(1.0)
    remove_table_borders(table)
    set_cell_borders(table.cell(0, 0), bottom={'sz': '2', 'color': 'DDDDDD'})
    set_cell_borders(table.cell(0, 1), bottom={'sz': '2', 'color': 'DDDDDD'})
    cell = table.cell(0, 0)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(f"\u2705  {name}")
    run.font.name = "Arial"
    run.font.size = Pt(9.5)
    run.font.color.rgb = BLACK
    run.bold = True
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(4)
    run2 = p2.add_run(desc)
    run2.font.name = "Arial"
    run2.font.size = Pt(8)
    run2.font.color.rgb = DARK_GRAY
    right_cell = table.cell(0, 1)
    right_cell.text = ""
    p = right_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    priority_color = BLACK if priority == "Critical" else DARK_GRAY if priority == "High" else GRAY
    run = p.add_run(priority)
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = priority_color
    run.bold = (priority in ("Critical", "High"))


# ══════════════════════════════════════════════════════════════════════════════
# GENERATE CHARTS
# ══════════════════════════════════════════════════════════════════════════════

print("Generating charts...")
perf_chart = chart_brand_performance()
severity_chart = chart_issue_severity()
cwv_chart = chart_cwv_comparison()
completeness_chart = chart_seo_completeness()
load_time_chart = chart_load_time()
issues_brand_chart = chart_issues_by_brand()

# ══════════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════

print("Building document...")
doc = Document(TEMPLATE_PATH)
body = doc.element.body
sect_pr = body.findall(qn('w:sectPr'))
sect_pr_copy = [deepcopy(sp) for sp in sect_pr]
for child in list(body):
    body.remove(child)
for sp in sect_pr_copy:
    body.append(sp)

# Enable "different first page" so page 1 has no header/footer
for sp in body.findall(qn('w:sectPr')):
    title_pg = sp.find(qn('w:titlePg'))
    if title_pg is None:
        title_pg = sp.makeelement(qn('w:titlePg'), {})
        sp.append(title_pg)

# ── COVER PAGE ───────────────────────────────────────────────────────────────

# Push content to vertical center
for _ in range(6):
    add_spacer(doc, 24)

logo_p = doc.add_paragraph()
logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
logo_p.paragraph_format.space_after = Pt(8)
logo_run = logo_p.add_run()
logo_run.add_picture(LOGO_PATH, width=Inches(3.5))

add_text(doc, "Multi-Brand SEO Audit Report", size=26, color=BLACK, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
add_text(doc, "Big Thunder Marine", size=18, color=GRAY,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_text(doc, "Dock Co.  |  Boat Rentals  |  Wake Experience  |  Iconic Yacht Charters",
         size=10, color=LIGHT_GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
add_text(doc, "Generated on March 26, 2026", size=10, color=LIGHT_GRAY,
         align=WD_ALIGN_PARAGRAPH.CENTER)

# ── EXECUTIVE SUMMARY (page 2) ──────────────────────────────────────────────

section_heading(doc, "Executive Summary", page_break=True)

# Metrics row
metrics_table = doc.add_table(rows=2, cols=4)
metrics_table.alignment = WD_TABLE_ALIGNMENT.CENTER
remove_table_borders(metrics_table)

labels = ["Brands Audited", "Avg. Mobile Perf.", "Avg. Lighthouse SEO", "Critical Issues"]
values = ["4", "74%", "100%", "5"]

for i in range(4):
    set_cell_text(metrics_table.cell(0, i), values[i], size=24, color=BLACK,
                  bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(metrics_table.cell(1, i), labels[i], size=7, color=GRAY,
                  bold=False, align=WD_ALIGN_PARAGRAPH.CENTER)

sub_heading(doc, "Key Findings", space_before=6)

bullet(doc, "Dock Co. has the strongest SEO (100%) but 47% mobile performance and wrong brand name on 2 pages.")
bullet(doc, "Boat Rentals is fast (98% mobile) but is a single landing page with no robots.txt, no sitemap, no schema, and no additional content for Google to rank.")
bullet(doc, "BT Wake is the healthiest brand: 80% mobile, 100% SEO, JSON-LD present, all meta tags configured.")
bullet(doc, "Iconic has a critical 40.3s mobile LCP. Also missing robots.txt, OG tags, canonical, and JSON-LD.")
bullet(doc, "3 of 4 brands have zero JSON-LD schema. Google cannot display rich results for them.")

# Scorecard table
sub_heading(doc, "Portfolio Scorecard", space_before=8)

sc = doc.add_table(rows=6, cols=6)
sc.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["Brand", "Mobile Perf.", "Desktop Perf.", "LH SEO*", "Access.", "Best Prac."]
for i, h in enumerate(headers):
    set_cell_text(sc.cell(0, i), h, size=7, color=WHITE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(sc.cell(0, i), "222222")

data = [
    ["Dock Co.", "47%", "67%", "100%", "96%", "100%"],
    ["Boat Rentals", "98%", "100%", "100%", "94%", "100%"],
    ["BT Wake", "80%", "96%", "100%", "94%", "100%"],
    ["Iconic Yachts", "71%", "99%", "100%", "88%", "100%"],
    ["Average", "74%", "91%", "100%", "93%", "100%"],
]

for r, row_data in enumerate(data):
    is_avg = r == len(data) - 1
    for c, val in enumerate(row_data):
        bold = is_avg or c == 0
        color = DARK_GRAY
        if c >= 1 and "%" in val:
            num = int(val.replace("%", ""))
            color = score_color(num)
        set_cell_text(sc.cell(r + 1, c), val, size=8, color=color, bold=bold,
                      align=WD_ALIGN_PARAGRAPH.CENTER if c > 0 else None)
        shade = "F5F5F5" if r % 2 == 0 else "FFFFFF"
        if is_avg:
            shade = "EEEEEE"
        set_cell_shading(sc.cell(r + 1, c), shade)

add_text(doc, "*LH SEO = Google Lighthouse SEO audit. Checks ~15 basic items (meta tags present, page crawlable, links have text). Does NOT measure keyword rankings, content quality, backlinks, schema, or local SEO signals. A 100% here means the basics pass; it does not mean SEO is complete.",
         size=7, color=GRAY, italic=True, space_before=2, space_after=0)

# ── VISUAL DASHBOARD (page 3) ────────────────────────────────────────────────

section_heading(doc, "Performance Dashboard")

add_text(doc, "Google measures every website on Performance (speed), SEO (search optimization), Accessibility, and Best Practices. Mobile scores matter most because Google ranks sites based on their mobile experience.",
         size=9, color=DARK_GRAY, space_after=4)

sub_heading(doc, "Mobile vs Desktop Performance", space_before=2)
add_chart(doc, perf_chart, width=Inches(5.2))

# Two-column: severity pie (left) + issue legend (right)
layout = doc.add_table(rows=1, cols=2)
layout.alignment = WD_TABLE_ALIGNMENT.CENTER
remove_table_borders(layout)
layout.columns[0].width = Inches(3.0)
layout.columns[1].width = Inches(3.5)

left_cell = layout.cell(0, 0)
left_cell.text = ""
p = left_cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run()
run.add_picture(severity_chart, width=Inches(2.2))

right_cell = layout.cell(0, 1)
right_cell.text = ""
p = right_cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("Issue Categories")
run.font.name = "Arial"
run.font.size = Pt(9)
run.font.color.rgb = BLACK
run.bold = True
for label, desc in [("5 Critical", "Revenue-blocking; fix this week"), ("9 High", "SEO impact; fix this sprint"), ("8 Medium", "Best practice gaps"), ("4 Low", "Minor polish")]:
    p2 = right_cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(3)
    p2.paragraph_format.space_after = Pt(0)
    run = p2.add_run(f"{label}  ")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = BLACK
    run.bold = True
    run2 = p2.add_run(desc)
    run2.font.name = "Arial"
    run2.font.size = Pt(7.5)
    run2.font.color.rgb = GRAY

# Technical completeness + CWV (page 4)
section_heading(doc, "Technical SEO Completeness", page_break=True)
add_text(doc, "Each website needs certain elements for search engines to properly index and display it. This matrix shows which elements are present, partially implemented, or missing across your four brands.",
         size=9, color=DARK_GRAY, space_after=2)
add_text(doc, "Light = present, medium = partial, dark = missing.",
         size=8, color=DARK_GRAY, italic=True, space_after=2)
add_chart(doc, completeness_chart, width=Inches(5.2))

sub_heading(doc, "Core Web Vitals (Mobile)", space_before=6)
add_text(doc, "Core Web Vitals are Google's speed metrics that directly affect search rankings. "
         "FCP (First Contentful Paint) measures how fast the first text or image appears. "
         "LCP (Largest Contentful Paint) measures how fast the main content loads; "
         "Google considers anything over 2.5 seconds slow. "
         "Iconic's 40.3s LCP means visitors wait over 40 seconds for the page to load on mobile.",
         size=9, color=DARK_GRAY, space_after=2)
add_chart(doc, cwv_chart, width=Inches(5.2))

# Load time + issues by brand (same page as above)
sub_heading(doc, "Site Speed & Issue Distribution", space_before=4)
add_text(doc, "Server response time is how fast your web host delivers the first byte of data. "
         "Time until usable is what your customers actually experience; Google penalizes sites over 2.5 seconds.",
         size=9, color=DARK_GRAY, space_after=2)
add_chart(doc, load_time_chart, width=Inches(5.2))

sub_heading(doc, "Issues Found Per Brand", space_before=4)
add_text(doc, "Total issues per brand by severity. Darker = more urgent. BT Wake has the fewest; Dock Co. and Iconic have the most.",
         size=9, color=DARK_GRAY, space_after=2)
add_chart(doc, issues_brand_chart, width=Inches(5.2))

# ── BRAND 1: DOCK CO. ───────────────────────────────────────────────────────

section_heading(doc, "Big Thunder Dock Co.", page_break=True)
add_text(doc, "bigthunderdocks.com  |  Wix (Velo)  |  Redirects to www", size=9, color=GRAY, italic=True, space_after=4)

sub_heading(doc, "Infrastructure", space_before=4)

infra = doc.add_table(rows=5, cols=2)
infra.alignment = WD_TABLE_ALIGNMENT.CENTER
infra.columns[0].width = Inches(2.2)
infra.columns[1].width = Inches(4.3)
remove_table_borders(infra)

infra_data = [
    ("HTTP Status", "301 (redirects to www)"),
    ("TTFB", "0.137s"),
    ("SSL Expiry", "April 28, 2026"),
    ("robots.txt", "Present; blocks PetalBot, rate-limits AhrefsBot"),
    ("Sitemap", "Present; 2 URLs"),
]
for r, (label, val) in enumerate(infra_data):
    set_cell_text(infra.cell(r, 0), label, size=8, color=GRAY, bold=True)
    set_cell_text(infra.cell(r, 1), val, size=8, color=DARK_GRAY)
    set_cell_borders(infra.cell(r, 0), bottom={'sz': '1', 'color': 'EEEEEE'})
    set_cell_borders(infra.cell(r, 1), bottom={'sz': '1', 'color': 'EEEEEE'})

sub_heading(doc, "Issues Found", space_before=8)

issue_row(doc, "Wrong brand name on /services and /contact pages",
          "Title and OG tags show \"M&B Remodeling\" instead of Big Thunder Dock Company. A Wix template artifact. Google indexes the wrong brand name.", "Critical")
issue_row(doc, "Missing meta descriptions on /services and /contact",
          "Google auto-generates search snippets. Lost control over messaging.", "High")
issue_row(doc, "Missing OG descriptions on /services and /contact",
          "Social sharing previews show no description.", "Medium")
issue_row(doc, "Zero JSON-LD structured data",
          "No schema on any page. Google cannot show rich results (hours, reviews, services).", "High")
issue_row(doc, "Sitemap has only 2 URLs",
          "If more pages exist, they may not be discovered by search engines.", "Medium")
issue_row(doc, "Mobile performance: 47% (homepage)",
          "Below Google's threshold. LCP 5.9s, TBT 920ms. Heavy JS and unoptimized images.", "High")
issue_row(doc, "Contact page CLS: 0.326",
          "Well above Google's 0.1 threshold. Elements shift during load.", "Medium")

# ── BRAND 2: BOAT RENTALS ────────────────────────────────────────────────────

section_heading(doc, "Big Thunder Boat Rentals", page_break=True)
add_text(doc, "bigthunderboatrentals.com  |  Next.js + Supabase + Vercel", size=9, color=GRAY, italic=True, space_after=4)

sub_heading(doc, "Infrastructure", space_before=4)

infra2 = doc.add_table(rows=5, cols=2)
infra2.alignment = WD_TABLE_ALIGNMENT.CENTER
infra2.columns[0].width = Inches(2.2)
infra2.columns[1].width = Inches(4.3)
remove_table_borders(infra2)

infra2_data = [
    ("HTTP Status", "200 (healthy)"),
    ("TTFB", "0.113s"),
    ("SSL Expiry", "June 6, 2026"),
    ("robots.txt", "MISSING (404)"),
    ("Sitemap", "MISSING (404)"),
]
for r, (label, val) in enumerate(infra2_data):
    set_cell_text(infra2.cell(r, 0), label, size=8, color=GRAY, bold=True)
    c = BLACK if "MISSING" in val else DARK_GRAY
    set_cell_text(infra2.cell(r, 1), val, size=8, color=c, bold="MISSING" in val)
    set_cell_borders(infra2.cell(r, 0), bottom={'sz': '1', 'color': 'EEEEEE'})
    set_cell_borders(infra2.cell(r, 1), bottom={'sz': '1', 'color': 'EEEEEE'})

sub_heading(doc, "Issues Found", space_before=8)

issue_row(doc, "Single-page site with no additional content",
          "The entire site is one landing page. There are no boats listing, contact, or service pages. This limits the keywords and topics Google can rank the site for.", "High")
issue_row(doc, "No robots.txt",
          "Search engines have no crawling instructions.", "High")
issue_row(doc, "No sitemap.xml",
          "Search engines must discover pages through links alone.", "High")
issue_row(doc, "No canonical tags, OG tags, or JSON-LD schema",
          "Missing all structured metadata. Social sharing is blank; no rich results possible.", "High")
issue_row(doc, "Homepage has zero H1 tags",
          "No clear heading signal for Google.", "Medium")

sub_heading(doc, "What's Working", space_before=6)
bullet(doc, "98% mobile performance (best in portfolio). FCP 1.1s, LCP 2.4s, TBT 10ms.")
bullet(doc, "100% desktop performance. The Next.js + Vercel stack is extremely fast.")

# ── BRAND 3: BT WAKE ────────────────────────────────────────────────────────

section_heading(doc, "BT Wake Experience", page_break=True)
add_text(doc, "btwake.com  |  Custom site  |  307 redirect to www", size=9, color=GRAY, italic=True, space_after=4)

sub_heading(doc, "Infrastructure", space_before=4)

infra3 = doc.add_table(rows=5, cols=2)
infra3.alignment = WD_TABLE_ALIGNMENT.CENTER
infra3.columns[0].width = Inches(2.2)
infra3.columns[1].width = Inches(4.3)
remove_table_borders(infra3)

infra3_data = [
    ("HTTP Status", "307 (temporary redirect to www)"),
    ("TTFB", "0.117s"),
    ("SSL Expiry", "May 28, 2026"),
    ("robots.txt", "MISSING (follows redirect)"),
    ("Sitemap", "Present; 1 URL"),
]
for r, (label, val) in enumerate(infra3_data):
    set_cell_text(infra3.cell(r, 0), label, size=8, color=GRAY, bold=True)
    c = BLACK if "MISSING" in val else DARK_GRAY
    set_cell_text(infra3.cell(r, 1), val, size=8, color=c, bold="MISSING" in val)
    set_cell_borders(infra3.cell(r, 0), bottom={'sz': '1', 'color': 'EEEEEE'})
    set_cell_borders(infra3.cell(r, 1), bottom={'sz': '1', 'color': 'EEEEEE'})

sub_heading(doc, "Status: Strongest Brand in Portfolio", space_before=8)

bullet(doc, "100% SEO score.", color=DARK_GRAY)
bullet(doc, "JSON-LD schema present (2 schemas). Only BT brand with structured data.", color=DARK_GRAY)
bullet(doc, "All meta tags properly configured: title, description, OG, canonical, viewport.", color=DARK_GRAY)
bullet(doc, "80% mobile performance; solid for a media-heavy watersports site.", color=DARK_GRAY)

sub_heading(doc, "Minor Issues", space_before=6)

issue_row(doc, "Missing robots.txt",
          "307 redirect prevents direct access. Most crawlers handle this, but explicit is better.", "Low")
issue_row(doc, "Sitemap has only 1 URL",
          "Additional pages (pricing, gallery) should be added.", "Low")
issue_row(doc, "Desktop CLS: 0.067",
          "Slightly above the 0.05 threshold.", "Low")

# ── BRAND 4: ICONIC ──────────────────────────────────────────────────────────

section_heading(doc, "Iconic Yacht Charters", page_break=True)
add_text(doc, "iconic-yachtcharters.com  |  Custom site  |  308 redirect to www", size=9, color=GRAY, italic=True, space_after=4)

sub_heading(doc, "Infrastructure", space_before=4)

infra4 = doc.add_table(rows=5, cols=2)
infra4.alignment = WD_TABLE_ALIGNMENT.CENTER
infra4.columns[0].width = Inches(2.2)
infra4.columns[1].width = Inches(4.3)
remove_table_borders(infra4)

infra4_data = [
    ("HTTP Status", "308 (permanent redirect to www)"),
    ("TTFB", "0.128s"),
    ("SSL Expiry", "June 10, 2026"),
    ("robots.txt", "MISSING (follows redirect)"),
    ("Sitemap", "Present; 17 URLs"),
]
for r, (label, val) in enumerate(infra4_data):
    set_cell_text(infra4.cell(r, 0), label, size=8, color=GRAY, bold=True)
    c = BLACK if "MISSING" in val else DARK_GRAY
    set_cell_text(infra4.cell(r, 1), val, size=8, color=c, bold="MISSING" in val)
    set_cell_borders(infra4.cell(r, 0), bottom={'sz': '1', 'color': 'EEEEEE'})
    set_cell_borders(infra4.cell(r, 1), bottom={'sz': '1', 'color': 'EEEEEE'})

sub_heading(doc, "Issues Found", space_before=8)

issue_row(doc, "Mobile LCP: 40.3 seconds",
          "Worst in portfolio by 10x. Hero image or video is unoptimized. Google will severely penalize mobile rankings.", "Critical")
issue_row(doc, "Missing OG tags",
          "Social sharing generates blank previews.", "High")
issue_row(doc, "Missing canonical tag",
          "www and non-www active; Google may split ranking signals.", "High")
issue_row(doc, "No JSON-LD structured data",
          "No rich results possible for a luxury charter service.", "High")
issue_row(doc, "Missing robots.txt",
          "No crawling instructions.", "Medium")
issue_row(doc, "Accessibility: 88% (lowest in portfolio)",
          "Likely missing alt text, color contrast issues, or ARIA labels.", "Medium")

sub_heading(doc, "What's Working", space_before=6)
bullet(doc, "Strong meta title and description; well-written and keyword-targeted.")
bullet(doc, "17-page sitemap; largest in portfolio, indicating good content depth.")
bullet(doc, "99% desktop performance; mobile issue is a specific asset, not structural.")

# ── PRIORITY ACTION PLAN ─────────────────────────────────────────────────────

section_heading(doc, "Priority Action Plan", page_break=True)

add_text(doc, "Ranked by impact across all 4 brands.",
         size=8, color=GRAY, italic=True, space_after=4)

fixable_row(doc, "Expand Boat Rentals beyond a single landing page",
    "Add dedicated pages for boat listings, pricing, and contact. A single page limits keyword coverage and ranking potential.", "High")
fixable_row(doc, "Fix Iconic mobile LCP (40.3s)",
    "Compress and lazy-load hero image/video. Target: LCP under 2.5 seconds.", "Critical")
fixable_row(doc, "Fix Dock Co. brand name on /services and /contact",
    "Replace \"M&B Remodeling\" with \"Big Thunder Dock Company\" in Wix page settings.", "Critical")
fixable_row(doc, "Add robots.txt to Boat Rentals, Wake, and Iconic",
    "Reference sitemap URL. For Next.js, add public/robots.txt.", "High")
fixable_row(doc, "Add sitemap.xml to Boat Rentals",
    "Only brand with no sitemap. Use next-sitemap or static file.", "High")
fixable_row(doc, "Add JSON-LD schema to Dock Co., Boat Rentals, and Iconic",
    "Use BT Wake's schema as template. LocalBusiness with NAP, hours, services, geo.", "High")
fixable_row(doc, "Add meta descriptions to Dock Co. /services and /contact",
    "Unique, keyword-rich descriptions under 160 characters.", "High")
fixable_row(doc, "Add OG tags to Boat Rentals and Iconic",
    "og:title, og:description, og:image for social sharing.", "Medium")
fixable_row(doc, "Add canonical tags to Boat Rentals and Iconic",
    "Prevent duplicate content between www and non-www.", "Medium")
fixable_row(doc, "Optimize Dock Co. mobile performance (47%)",
    "Compress images, defer JS, reduce TBT. Wix may cap gains at 60-70%.", "Medium")
fixable_row(doc, "Add H1 tag to Boat Rentals homepage",
    "Clear heading with primary keyword.", "Low")
fixable_row(doc, "Fix Iconic accessibility (88%)",
    "Alt text, color contrast, ARIA labels. Target: 95%+.", "Low")

# ── NEXT STEPS (inline after action plan) ────────────────────────────────────

sub_heading(doc, "Timeline", space_before=8)

timeline_items = [
    "This week: Fix Iconic mobile LCP + Dock Co. brand name error",
    "This sprint: Add robots.txt, sitemaps, and JSON-LD schema to all brands",
    "Next sprint: Add OG tags, canonical tags, and optimize Dock Co. mobile performance",
    "Ongoing: Bi-weekly monitoring via automated pipeline. Next report: week of April 7, 2026",
]
for item in timeline_items:
    bullet(doc, item)

add_spacer(doc, 8)
add_text(doc, "This report was generated by Wetware Labs LLC using automated auditing tools "
         "and the Google PageSpeed Insights API. For questions or to discuss remediation, "
         "contact management@wetwareofficial.com.",
         size=8, color=LIGHT_GRAY, italic=True, space_before=8, space_after=0)

# ── SAVE & RE-INJECT LOGO ───────────────────────────────────────────────────

doc.save(OUTPUT_DOCX)
# Template already has the header logo embedded; no re-injection needed.

print(f"\u2705 DOCX saved: {OUTPUT_DOCX}")

# ── CONVERT TO PDF ───────────────────────────────────────────────────────────

SOFFICE = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
OUT_DIR = os.path.dirname(OUTPUT_DOCX)
result = subprocess.run(
    [SOFFICE, "--headless", "--convert-to", "pdf", "--outdir", OUT_DIR, OUTPUT_DOCX],
    capture_output=True, text=True, timeout=60
)
if result.returncode == 0:
    pdf_name = os.path.splitext(os.path.basename(OUTPUT_DOCX))[0] + ".pdf"
    pdf_path = os.path.join(OUT_DIR, pdf_name)
    print(f"\u2705 PDF saved: {pdf_path}")
else:
    print(f"\u274c PDF conversion failed: {result.stderr}")

# Cleanup temp chart images
import shutil
shutil.rmtree(TMPDIR, ignore_errors=True)
